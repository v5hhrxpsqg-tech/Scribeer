import os
from flask import Flask, render_template, session, redirect, url_for, request, flash, Response, send_file
from io import BytesIO
from docx import Document
from fpdf import FPDF
from dotenv import load_dotenv
from supabase import create_client
from deepgram import DeepgramClient
from datetime import datetime
import stripe

# 1. Gegevens laden uit .env
load_dotenv()

app = Flask(__name__)
app.secret_key = os.getenv("FLASK_SECRET_KEY")
app.config['MAX_CONTENT_LENGTH'] = 200 * 1024 * 1024  # Max 200MB upload

# 2. Supabase koppelen
supabase = create_client(os.getenv("SUPABASE_URL"), os.getenv("SUPABASE_KEY"))

# 3. Deepgram koppelen
deepgram = DeepgramClient(api_key=os.getenv("DEEPGRAM_API_KEY"))

# 4. Stripe koppelen
stripe.api_key = os.getenv("STRIPE_SECRET_KEY")

# Credit pakketten (minuten, prijs, Stripe Payment Link)
CREDIT_PACKAGES = {
    'small': {'minutes': 100, 'price': 499, 'name': '100 minuten', 'stripe_link': 'https://buy.stripe.com/test_6oUcN5e65fuz2uYdWveZ200'},
    'medium': {'minutes': 300, 'price': 1199, 'name': '300 minuten', 'stripe_link': 'https://buy.stripe.com/test_8x2eVdd2196bb1u3hReZ201'},
    'large': {'minutes': 500, 'price': 1599, 'name': '500 minuten', 'stripe_link': 'https://buy.stripe.com/test_00w6oH4vv0zF3z25pZeZ202'},
}

# =================================================================
# HELPER FUNCTIES
# =================================================================
def get_or_create_user_credits(user_email: str) -> dict:
    """Haalt user credits op, of maakt nieuwe user aan met 50 MB."""
    result = supabase.table('user_credits').select('*').eq('email', user_email).execute()

    if result.data:
        return result.data[0]
    else:
        new_user = supabase.table('user_credits').insert({
            'email': user_email,
            'credits_remaining_mb': 50.00
        }).execute()
        return new_user.data[0]

def deduct_credits(user_email: str, mb_used: float) -> bool:
    """Trekt credits af. Returns True als gelukt."""
    current = get_or_create_user_credits(user_email)
    new_balance = float(current['credits_remaining_mb']) - mb_used

    if new_balance < 0:
        return False

    supabase.table('user_credits').update({
        'credits_remaining_mb': new_balance
    }).eq('email', user_email).execute()
    return True

def add_credits(user_email: str, minutes_to_add: float):
    """Voegt credits toe aan gebruiker."""
    current = get_or_create_user_credits(user_email)
    new_balance = float(current['credits_remaining_mb']) + minutes_to_add

    supabase.table('user_credits').update({
        'credits_remaining_mb': new_balance
    }).eq('email', user_email).execute()

def save_transcription(user_id: str, user_email: str, filename: str, transcript: str, duration_minutes: float):
    """Slaat transcriptie op in de database."""
    supabase.table('transcriptions').insert({
        'user_id': user_id,
        'user_email': user_email,
        'filename': filename,
        'transcript': transcript,
        'duration_minutes': duration_minutes,
        'created_at': datetime.utcnow().isoformat()
    }).execute()

def get_user_transcriptions(user_email: str):
    """Haalt alle transcripties van een gebruiker op."""
    result = supabase.table('transcriptions').select('*').eq('user_email', user_email).order('created_at', desc=True).execute()
    return result.data

def transcribe_audio(audio_data: bytes, filename: str) -> dict:
    """Transcribeert audio met Deepgram. Returns dict met transcript en metadata."""
    # Roep Deepgram API aan met keyword arguments
    response = deepgram.listen.v1.media.transcribe_file(
        request=audio_data,
        model="nova-2",
        language="nl",
        smart_format=True,
        diarize=True,
        paragraphs=True,
        utterances=True
    )

    # Haal resultaten uit response
    result = response.results
    transcript_text = result.channels[0].alternatives[0].transcript

    # Haal duur op in minuten
    duration_seconds = float(response.metadata.duration)
    duration_minutes = duration_seconds / 60

    # Formatteer met sprekers als beschikbaar
    paragraphs = result.channels[0].alternatives[0].paragraphs
    if paragraphs and paragraphs.paragraphs:
        formatted_text = ""
        for para in paragraphs.paragraphs:
            speaker = f"Spreker {para.speaker + 1}" if hasattr(para, 'speaker') else "Spreker"
            sentences = " ".join([s.text for s in para.sentences])
            formatted_text += f"**{speaker}:** {sentences}\n\n"
        transcript_text = formatted_text

    return {
        'transcript': transcript_text,
        'duration_minutes': duration_minutes
    }

# =================================================================
# ROUTES
# =================================================================

@app.route('/')
def home():
    """Hoofdpagina - redirect naar dashboard of login"""
    if 'user' in session:
        return redirect(url_for('dashboard'))
    return redirect(url_for('login'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    """Login pagina met magic link"""
    if 'user' in session:
        return redirect(url_for('dashboard'))

    if request.method == 'POST':
        email = request.form.get('email')
        if email:
            try:
                supabase.auth.sign_in_with_otp({
                    "email": email,
                    "options": {
                        "email_redirect_to": "http://127.0.0.1:5000/callback"
                    }
                })
                flash('Check je email voor de login link!', 'success')
            except Exception as e:
                flash(f'Fout bij verzenden: {str(e)}', 'error')
        else:
            flash('Vul een email adres in', 'error')

    return render_template('login.html')

@app.route('/callback')
def callback():
    """Toont de callback pagina die de hash fragment uitleest via JavaScript"""
    return render_template('callback.html')

@app.route('/auth/verify')
def auth_verify():
    """Verwerkt de token nadat JavaScript hem heeft doorgestuurd"""
    access_token = request.args.get('access_token')

    if access_token:
        try:
            user_response = supabase.auth.get_user(access_token)

            if user_response and user_response.user:
                session['user'] = {
                    'id': user_response.user.id,
                    'email': user_response.user.email
                }
                session['access_token'] = access_token
                flash('Succesvol ingelogd!', 'success')
                return redirect(url_for('dashboard'))
        except Exception as e:
            flash(f'Login fout: {str(e)}', 'error')

    flash('Login mislukt. Probeer opnieuw.', 'error')
    return redirect(url_for('login'))

@app.route('/dashboard')
def dashboard():
    """Dashboard - alleen voor ingelogde gebruikers"""
    if 'user' not in session:
        flash('Je moet eerst inloggen', 'error')
        return redirect(url_for('login'))

    user = session['user']

    try:
        user_credits = get_or_create_user_credits(user['email'])
        credits = float(user_credits['credits_remaining_mb'])
    except Exception as e:
        credits = 0
        flash(f'Kon credits niet laden: {str(e)}', 'error')

    # Haal recente transcripties op
    try:
        transcriptions = get_user_transcriptions(user['email'])
    except:
        transcriptions = []

    return render_template('dashboard.html', user=user, credits=credits, transcriptions=transcriptions)

@app.route('/upload', methods=['POST'])
def upload():
    """Verwerkt audio upload en start transcriptie"""
    if 'user' not in session:
        flash('Je moet eerst inloggen', 'error')
        return redirect(url_for('login'))

    user = session['user']

    # Check of er een bestand is
    if 'audio' not in request.files:
        flash('Geen bestand geselecteerd', 'error')
        return redirect(url_for('dashboard'))

    file = request.files['audio']
    if file.filename == '':
        flash('Geen bestand geselecteerd', 'error')
        return redirect(url_for('dashboard'))

    # Check bestandstype
    allowed_extensions = {'mp3', 'wav', 'm4a', 'ogg', 'flac'}
    ext = file.filename.lower().split('.')[-1]
    if ext not in allowed_extensions:
        flash('Ongeldig bestandstype. Gebruik MP3, WAV, M4A, OGG of FLAC.', 'error')
        return redirect(url_for('dashboard'))

    # Lees bestand
    audio_data = file.read()
    file_mb = len(audio_data) / (1024 * 1024)

    # Check credits
    try:
        user_credits = get_or_create_user_credits(user['email'])
        credits = float(user_credits['credits_remaining_mb'])

        if credits < file_mb:
            flash(f'Onvoldoende credits. Je hebt {credits:.1f} minuten, maar dit bestand is {file_mb:.1f} MB.', 'error')
            return redirect(url_for('dashboard'))
    except Exception as e:
        flash(f'Kon credits niet controleren: {str(e)}', 'error')
        return redirect(url_for('dashboard'))

    # Transcribeer met Deepgram
    try:
        result = transcribe_audio(audio_data, file.filename)
        transcript = result['transcript']
        duration = result['duration_minutes']

        # Trek credits af (op basis van audio duur, niet bestandsgrootte)
        deduct_credits(user['email'], duration)

        # Sla op in database
        save_transcription(
            user_id=user['id'],
            user_email=user['email'],
            filename=file.filename,
            transcript=transcript,
            duration_minutes=duration
        )

        flash(f'Transcriptie voltooid! ({duration:.1f} minuten verwerkt)', 'success')

    except Exception as e:
        flash(f'Transcriptie mislukt: {str(e)}', 'error')

    return redirect(url_for('dashboard'))

@app.route('/transcription/<int:transcription_id>')
def view_transcription(transcription_id):
    """Bekijk een specifieke transcriptie"""
    if 'user' not in session:
        flash('Je moet eerst inloggen', 'error')
        return redirect(url_for('login'))

    user = session['user']

    # Haal transcriptie op
    result = supabase.table('transcriptions').select('*').eq('id', transcription_id).eq('user_email', user['email']).execute()

    if not result.data:
        flash('Transcriptie niet gevonden', 'error')
        return redirect(url_for('dashboard'))

    transcription = result.data[0]
    return render_template('transcription.html', user=user, transcription=transcription)

@app.route('/download/<int:transcription_id>/<format>')
def download_transcription(transcription_id, format):
    """Download transcriptie als Word of PDF"""
    if 'user' not in session:
        flash('Je moet eerst inloggen', 'error')
        return redirect(url_for('login'))

    user = session['user']

    # Haal transcriptie op
    result = supabase.table('transcriptions').select('*').eq('id', transcription_id).eq('user_email', user['email']).execute()

    if not result.data:
        flash('Transcriptie niet gevonden', 'error')
        return redirect(url_for('dashboard'))

    transcription = result.data[0]
    transcript_text = transcription['transcript']
    filename_base = transcription['filename'].rsplit('.', 1)[0]

    # Verwijder markdown formatting voor schone tekst
    clean_text = transcript_text.replace('**', '')

    if format == 'docx':
        # Maak Word document
        doc = Document()
        doc.add_heading(f'Transcriptie: {transcription["filename"]}', 0)
        doc.add_paragraph(f'Datum: {transcription["created_at"][:10]}')
        doc.add_paragraph(f'Duur: {transcription["duration_minutes"]:.1f} minuten')
        doc.add_paragraph('')

        # Voeg transcriptie toe per paragraaf
        for para in clean_text.split('\n\n'):
            if para.strip():
                doc.add_paragraph(para.strip())

        # Sla op in memory
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return Response(
            buffer.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={'Content-Disposition': f'attachment; filename={filename_base}_transcriptie.docx'}
        )

    elif format == 'pdf':
        try:
            # Maak PDF document
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)

            # Titel
            pdf.set_font('Helvetica', 'B', 16)
            pdf.cell(0, 10, f'Transcriptie: {transcription["filename"]}', new_x='LMARGIN', new_y='NEXT')

            # Metadata
            pdf.set_font('Helvetica', '', 10)
            pdf.cell(0, 6, f'Datum: {transcription["created_at"][:10]}', new_x='LMARGIN', new_y='NEXT')
            pdf.cell(0, 6, f'Duur: {transcription["duration_minutes"]:.1f} minuten', new_x='LMARGIN', new_y='NEXT')
            pdf.ln(10)

            # Transcriptie tekst
            pdf.set_font('Helvetica', '', 11)
            # Encode tekst voor PDF compatibiliteit
            safe_text = clean_text.encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 6, safe_text)

            # Output naar BytesIO buffer
            pdf_buffer = BytesIO()
            pdf_buffer.write(pdf.output())
            pdf_buffer.seek(0)

            return send_file(
                pdf_buffer,
                mimetype='application/pdf',
                as_attachment=True,
                download_name=f'{filename_base}_transcriptie.pdf'
            )
        except Exception as e:
            print(f"PDF error: {str(e)}")
            flash(f'PDF generatie mislukt: {str(e)}', 'error')
            return redirect(url_for('view_transcription', transcription_id=transcription_id))

    else:
        flash('Ongeldig formaat', 'error')
        return redirect(url_for('view_transcription', transcription_id=transcription_id))

@app.route('/pricing')
def pricing():
    """Pricing pagina met credit pakketten"""
    if 'user' not in session:
        flash('Je moet eerst inloggen', 'error')
        return redirect(url_for('login'))

    user = session['user']
    return render_template('pricing.html', user=user, packages=CREDIT_PACKAGES)

@app.route('/checkout/<package_id>')
def checkout(package_id):
    """Redirect naar Stripe Payment Link"""
    if 'user' not in session:
        flash('Je moet eerst inloggen', 'error')
        return redirect(url_for('login'))

    if package_id not in CREDIT_PACKAGES:
        flash('Ongeldig pakket', 'error')
        return redirect(url_for('pricing'))

    package = CREDIT_PACKAGES[package_id]
    user = session['user']

    # Redirect naar Stripe Payment Link met client_reference_id voor identificatie
    stripe_url = f"{package['stripe_link']}?client_reference_id={user['email']}&prefilled_email={user['email']}"
    return redirect(stripe_url)

@app.route('/payment/success')
def payment_success():
    """Bedankpagina na succesvolle betaling - credits worden via Supabase webhook toegevoegd"""
    flash('Bedankt voor je aankoop! Je credits worden binnen enkele seconden toegevoegd.', 'success')
    return redirect(url_for('dashboard'))

@app.route('/logout')
def logout():
    """Uitloggen - verwijder sessie"""
    session.clear()
    flash('Je bent uitgelogd', 'success')
    return redirect(url_for('login'))

if __name__ == '__main__':
    app.run(debug=True, threaded=True)
